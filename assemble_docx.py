import docx
from docx.shared import Inches
import os

doc = docx.Document()
doc.add_heading('MediCart - Detailed Project Diagrams (from BlackBook)', 0)

titles = [
    "1. Data Flow Diagram — Level 0 (Context Diagram)",
    "2. Data Flow Diagram — Level 1",
    "3. Detailed Data Flow Diagram — 3.0 Order Processing",
    "4. Gantt Chart — Project Development Timeline",
    "5. PERT Chart — Project Task Dependencies",
    "6. ER Diagram (Entity Relationship Diagram)",
    "7. Use Case Diagram",
    "8. Activity Diagram — Doctor Order Flow",
    "9. Sequence Diagram — Order & Payment Flow",
    "10. State Transition Diagram",
    "11. Class Diagram",
    "12. Logic Flow Diagram — placeOrder() Algorithm",
    "13. Collaboration Diagram",
    "14. System Architecture Diagram"
]

for i, title in enumerate(titles, 1):
    doc.add_heading(title, level=2)
    img_path = f"extracted_diagrams_png-{i}.png"
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
    else:
        doc.add_paragraph(f"[Warning: Diagram Image {img_path} not found]")
    doc.add_page_break()

doc.save('MediCart_Diagrams_Final.docx')
print("Successfully created MediCart_Diagrams_Final.docx with 14 diagrams.")
